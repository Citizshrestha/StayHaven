from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('10-WEEK DAILY ACTIVITY LOG', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph('Hotel Booking Order Management System Development')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Timeline
timeline = doc.add_paragraph('Timeline: November 29, 2025 – February 6, 2026 | Work Schedule: 6 days per week')
timeline.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Week 1-2
doc.add_heading('WEEK 1-2: FOUNDATION & INITIAL SETUP (Nov 29 – Dec 12)', level=2)

table1 = doc.add_table(rows=13, cols=3)
table1.style = 'Light Grid Accent 1'

hdr_cells = table1.rows[0].cells
hdr_cells[0].text = 'Date'
hdr_cells[1].text = 'Activity Summary'
hdr_cells[2].text = 'Related Component(s)'

data1 = [
    ('2025-11-29 (Sat)', 'Project kickoff: reviewed requirements and architecture for Hotel admin panel. Set up development environment.', 'Hoteladmin'),
    ('2025-12-01 (Mon)', 'Created Hoteladmin UI skeleton with hotel listing page layout and navigation structure.', 'Hoteladmin / Frontend'),
    ('2025-12-02 (Tue)', 'Implemented search and filter functionality in Hoteladmin listings.', 'Hoteladmin / Frontend'),
    ('2025-12-03 (Wed)', 'Integrated Hoteladmin listing with backend API endpoints. Debugged CORS.', 'Hoteladmin / Routes / API'),
    ('2025-12-04 (Thu)', 'Code review: Hoteladmin components. Updated styling for consistency.', 'Hoteladmin / Code Quality'),
    ('2025-12-05 (Fri)', 'Wrote unit tests for Hoteladmin utilities and improved error handling.', 'Hoteladmin / Testing'),
    ('2025-12-06 (Sat)', 'Deployed Hoteladmin to staging environment. Executed smoke tests.', 'Hoteladmin / Deployment'),
    ('2025-12-08 (Mon)', 'Initialized Superadmin dashboard project structure and created dashboard page.', 'Superadmin'),
    ('2025-12-09 (Tue)', 'Implemented role-based access control UI in Superadmin with auth checks.', 'Superadmin / Auth'),
    ('2025-12-10 (Wed)', 'Built user management table with pagination, sorting, and filtering.', 'Superadmin'),
    ('2025-12-11 (Thu)', 'Connected Superadmin user actions to backend API routes. Tested CRUD.', 'Superadmin / Routes'),
    ('2025-12-12 (Fri)', 'Fixed Superadmin user role editing edge cases. Added validation.', 'Superadmin'),
]

for i, (date, activity, component) in enumerate(data1, 1):
    row_cells = table1.rows[i].cells
    row_cells[0].text = date
    row_cells[1].text = activity
    row_cells[2].text = component

doc.add_paragraph()

# Week 3-4
doc.add_heading('WEEK 3-4: FOOTER & BACKEND AUTHENTICATION (Dec 13 – Dec 26)', level=2)

table2 = doc.add_table(rows=12, cols=3)
table2.style = 'Light Grid Accent 1'

hdr_cells = table2.rows[0].cells
hdr_cells[0].text = 'Date'
hdr_cells[1].text = 'Activity Summary'
hdr_cells[2].text = 'Related Component(s)'

data2 = [
    ('2025-12-13 (Sat)', 'Code review: Superadmin components finalized. Deployed to test environment.', 'Superadmin'),
    ('2025-12-15 (Mon)', 'Implemented footer section with responsive layout and navigation links.', 'Footer'),
    ('2025-12-16 (Tue)', 'Finalized footer styling and accessibility features. Added social icons.', 'Footer'),
    ('2025-12-17 (Wed)', 'Integrated footer into shared layout across all pages.', 'Footer'),
    ('2025-12-18 (Thu)', 'Fixed footer overlap on mobile screens. Tested responsiveness.', 'Footer'),
    ('2025-12-19 (Fri)', 'Reviewed footer copy and added analytics tracking hooks.', 'Footer'),
    ('2025-12-20 (Sat)', 'Verified footer deployment. Reviewed auth routes and token refresh.', 'Footer / Auth'),
    ('2025-12-22 (Mon)', 'Implemented JWT token refresh in backend auth routes with validation.', 'Auth / Backend'),
    ('2025-12-23 (Tue)', 'Implemented role-based middleware in routes. Tested protected routes.', 'Routes / Auth'),
    ('2025-12-24 (Wed)', 'Debugged login flow. Fixed JWT expiry handling issues.', 'Auth'),
    ('2025-12-25 (Thu)', 'Wrote integration tests for auth endpoints. Ran CI/CD pipeline.', 'Auth / Testing'),
    ('2025-12-26 (Fri)', 'Fixed routes bug causing 500 errors. Added guard clauses.', 'Routes'),
]

for i, (date, activity, component) in enumerate(data2, 1):
    row_cells = table2.rows[i].cells
    row_cells[0].text = date
    row_cells[1].text = activity
    row_cells[2].text = component

doc.add_paragraph()

# Week 5-6
doc.add_heading('WEEK 5-6: KITCHEN DASHBOARD & MONITORING (Dec 27 – Jan 9)', level=2)

table3 = doc.add_table(rows=12, cols=3)
table3.style = 'Light Grid Accent 1'

hdr_cells = table3.rows[0].cells
hdr_cells[0].text = 'Date'
hdr_cells[1].text = 'Activity Summary'
hdr_cells[2].text = 'Related Component(s)'

data3 = [
    ('2025-12-27 (Sat)', 'Started monitoring pages in userguestfolder. Identified UI issues.', 'UserGuestFolder'),
    ('2025-12-29 (Mon)', 'Thorough review of userguestfolder pages. Reported bugs to team.', 'UserGuestFolder'),
    ('2025-12-30 (Tue)', 'Assisted with bug fixes in userguestfolder. Validated fixes.', 'UserGuestFolder'),
    ('2025-12-31 (Wed)', 'End-of-year health checks: verified auth flows and page integrity.', 'Auth / UserGuestFolder'),
    ('2026-01-01 (Thu)', 'Post-holiday verification: ensured auth flows operational.', 'Auth'),
    ('2026-01-02 (Fri)', 'Started monitoring KitchenDashboard. Captured performance metrics.', 'KitchenDashboard'),
    ('2026-01-03 (Sat)', 'Implemented KitchenDashboard widget for live order updates with WebSocket.', 'KitchenDashboard'),
    ('2026-01-05 (Mon)', 'Fixed KitchenDashboard socket reconnection issues. Improved resilience.', 'KitchenDashboard'),
    ('2026-01-06 (Tue)', 'Reviewed footer consistency in KitchenDashboard. Refined UI/UX.', 'KitchenDashboard / Footer'),
    ('2026-01-07 (Wed)', 'Paired with kitchen team to refine order prioritization logic.', 'KitchenDashboard'),
    ('2026-01-08 (Thu)', 'Documented KitchenDashboard API endpoints and data contracts.', 'KitchenDashboard'),
    ('2026-01-09 (Fri)', 'Refactored Hoteladmin state management. Improved performance.', 'Hoteladmin'),
]

for i, (date, activity, component) in enumerate(data3, 1):
    row_cells = table3.rows[i].cells
    row_cells[0].text = date
    row_cells[1].text = activity
    row_cells[2].text = component

doc.add_paragraph()

# Week 7
doc.add_heading('WEEK 7: MARKET RESEARCH & SUPERADMIN ENHANCEMENTS (Jan 10 – Jan 16)', level=2)

table4 = doc.add_table(rows=8, cols=3)
table4.style = 'Light Grid Accent 1'

hdr_cells = table4.rows[0].cells
hdr_cells[0].text = 'Date'
hdr_cells[1].text = 'Activity Summary'
hdr_cells[2].text = 'Related Component(s)'

data4 = [
    ('2026-01-10 (Sat)', 'Fixed Hoteladmin booking edge cases discovered in staging.', 'Hoteladmin'),
    ('2026-01-12 (Mon)', 'Implemented confirmation modal flows in Hoteladmin.', 'Hoteladmin'),
    ('2026-01-13 (Tue)', 'FIELD MARKET RESEARCH - DAY 1: On-site hotel visits for guest booking flows and pain points analysis.', 'Market Research'),
    ('2026-01-14 (Wed)', 'FIELD MARKET RESEARCH - DAY 2: Visited kitchens, observed staff workflows, identified KitchenDashboard optimization opportunities.', 'Market Research'),
    ('2026-01-15 (Thu)', 'Documented market research findings and created enhancement proposal.', 'Market Research'),
    ('2026-01-16 (Fri)', 'Worked on Superadmin analytics tile with data fetch retries.', 'Superadmin'),
    ('2026-01-17 (Sat)', 'Code review: Hoteladmin PR merged. Monitored post-release.', 'Hoteladmin'),
]

for i, (date, activity, component) in enumerate(data4, 1):
    row_cells = table4.rows[i].cells
    row_cells[0].text = date
    row_cells[1].text = activity
    row_cells[2].text = component

doc.add_paragraph()

# Week 8
doc.add_heading('WEEK 8: ADVANCED FEATURES & INTEGRATION (Jan 19 – Jan 29)', level=2)

table5 = doc.add_table(rows=11, cols=3)
table5.style = 'Light Grid Accent 1'

hdr_cells = table5.rows[0].cells
hdr_cells[0].text = 'Date'
hdr_cells[1].text = 'Activity Summary'
hdr_cells[2].text = 'Related Component(s)'

data5 = [
    ('2026-01-19 (Mon)', 'Added export-to-CSV functionality in Superadmin user reports.', 'Superadmin'),
    ('2026-01-20 (Tue)', 'Fixed Superadmin export permissions and large-file handling.', 'Superadmin'),
    ('2026-01-21 (Wed)', 'Reviewed auth logs to diagnose Superadmin export errors. Fixed token validation.', 'Auth / Superadmin'),
    ('2026-01-22 (Thu)', 'Polished footer: added dynamic year and newsletter signup form.', 'Footer'),
    ('2026-01-23 (Fri)', 'Accessibility audit of footer. Applied WCAG compliance fixes.', 'Footer'),
    ('2026-01-24 (Sat)', 'Created footer contact form endpoints. Implemented API integration.', 'Routes / Footer'),
    ('2026-01-26 (Mon)', 'Monitored userguestfolder page load times. Reported slow assets.', 'UserGuestFolder'),
    ('2026-01-27 (Tue)', 'Implemented lazy-loading for heavy assets in userguestfolder.', 'UserGuestFolder'),
    ('2026-01-28 (Wed)', 'Validated lazy-loading fixes. Cross-browser compatibility tests.', 'UserGuestFolder'),
    ('2026-01-29 (Thu)', 'Improved KitchenDashboard order filter performance and sorting.', 'KitchenDashboard'),
]

for i, (date, activity, component) in enumerate(data5, 1):
    row_cells = table5.rows[i].cells
    row_cells[0].text = date
    row_cells[1].text = activity
    row_cells[2].text = component

doc.add_paragraph()

# Week 9-10
doc.add_heading('WEEK 9-10: TESTING PHASE & FINAL INTEGRATION (Jan 30 – Feb 6)', level=2)

table6 = doc.add_table(rows=8, cols=3)
table6.style = 'Light Grid Accent 1'

hdr_cells = table6.rows[0].cells
hdr_cells[0].text = 'Date'
hdr_cells[1].text = 'Activity Summary'
hdr_cells[2].text = 'Related Component(s)'

data6 = [
    ('2026-01-30 (Fri)', 'Added retry/backoff mechanism for KitchenDashboard API calls.', 'KitchenDashboard'),
    ('2026-01-31 (Sat)', 'Implemented audit log view in Superadmin with backend integration.', 'Superadmin'),
    ('2026-02-02 (Mon)', 'Implemented role-based visibility for Superadmin audit logs.', 'Superadmin'),
    ('2026-02-03 (Tue)', 'Implemented inline editing for Hoteladmin room details.', 'Hoteladmin'),
    ('2026-02-04 (Wed)', 'Fixed inline edit validation in Hoteladmin with error recovery.', 'Hoteladmin'),
    ('2026-02-05 (Thu)', 'Comprehensive integration tests across Hoteladmin and Superadmin.', 'Hoteladmin / Superadmin'),
    ('2026-02-06 (Fri)', 'Tightened auth token refresh flow and security. PROJECT NOW IN TESTING PHASE.', 'Auth / Testing Phase'),
]

for i, (date, activity, component) in enumerate(data6, 1):
    row_cells = table6.rows[i].cells
    row_cells[0].text = date
    row_cells[1].text = activity
    row_cells[2].text = component

doc.add_paragraph()
doc.add_paragraph()

# Summary
summary = doc.add_heading('PROJECT STATUS: TESTING PHASE ACTIVE', level=2)
summary_text = doc.add_paragraph('All core components (Hoteladmin, Superadmin, Footer, KitchenDashboard, UserGuestFolder) implemented and integrated. Comprehensive QA underway. Project continues with full testing validation.')

# Save document
doc.save('10-Week-Activity-Log.docx')
print("Document created successfully: 10-Week-Activity-Log.docx")
